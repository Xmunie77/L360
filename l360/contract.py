"""Pre-filled tutor Services Agreement (.docx).

Generates the Foundation's "Contracts of Service — Tutor Agreement" with the
tutor's details merged in from their submitted onboarding form: name, ID
number, residential address, term start date, account holder and IBAN. The
clause text mirrors the template in Drive (HR → Educators →
"Contracts of Service - Tutor Agreement .docx") — edit wording there AND
here together.

Deliberately left blank for hand-completion at signing (this is a public
repo, so nothing personal or commercially sensitive is embedded):
- the founders' ID numbers in the parties block,
- the hourly rates in clause 5.2,
- the foundation email address the tutor will use (clause 3.5),
- the tutor's bank name / branch (not collected on the onboarding form).

The output is a starting point for the written agreement, not the executed
contract — it still gets reviewed, completed and signed on paper.
"""
from __future__ import annotations

import io
from datetime import date, datetime

from docx import Document
from docx.shared import Pt

_BLANK = "________________"


def _para(doc, text: str, *, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def build_contract(answers: dict, *, today: date | None = None) -> bytes:
    """answers = EducatorOnboardingForm.answers (validated submission)."""
    if today is None:
        # Malta calendar date — the server clock is UTC (see booking_logic.
        # local_today; not imported to keep this module docx-only/pure).
        from zoneinfo import ZoneInfo

        from l360.config import TIMEZONE

        today = datetime.now(ZoneInfo(TIMEZONE)).date()
    name = answers.get("full_legal_name") or _BLANK
    id_number = answers.get("id_passport_number") or _BLANK
    address = answers.get("residential_address") or _BLANK
    start = answers.get("preferred_start_date") or _BLANK
    account_holder = answers.get("bank_account_holder") or name
    iban = answers.get("iban") or _BLANK

    doc = Document()

    _para(doc, "SERVICES AGREEMENT", bold=True, size=14)
    _para(doc, (
        f"THIS SERVICES AGREEMENT (hereinafter referred to as the “Agreement”) is made as of "
        f"{today.strftime('%d %B %Y')}, by and between:"
    ))
    _para(doc, (
        f"LEARNING360° FOUNDATION, c/o Francesca Diacono holder of Maltese ID Number {_BLANK}, "
        f"Samantha Pace holder of Maltese ID Number {_BLANK} and Justine Balani holder of Maltese ID "
        f"Number {_BLANK} (hereinafter referred to as “the Foundation”)"
    ))
    _para(doc, "AND")
    _para(doc, (
        f"Ms/Mr {name}, holder of ID Number {id_number} and residing at the following address: "
        f"{address} (hereinafter referred to as the “Tutor”)."
    ), bold=True)
    _para(doc, "The Foundation and the Tutor are sometimes referred to herein as a “Party” and both of them together as the “Parties”.")

    _para(doc, "WHEREAS", bold=True)
    for t in (
        "A. The Foundation wishes to engage the services of the Tutor to provide individualized or group tuition services to the Foundation's clientele;",
        "B. Tutor desires to contract with the Foundation for the purposes of providing individualized tuition services to the Foundation in Malta;",
        "C. The Tutor has the required specialized skills and knowledge that are required by the Foundation and is fully informed of all aspects of the services required to be performed and agrees to accept the Foundation's offer to provide the services upon the terms and conditions contained in this Agreement; and",
        "D. The Parties agree that this Agreement constitutes the entire agreement between them with respect to the services to be rendered hereunder and supersedes all other agreements in connection with said services.",
    ):
        _para(doc, t)
    _para(doc, "NOW THEREFORE, it is hereby agreed as follows:", bold=True)

    sections: list[tuple[str, list[str]]] = [
        ("1. GENERAL", [
            "1.1 Words importing a gender shall include each other gender. Words in the singular include the plural, and words in the plural include the singular.",
            "1.2 The Parties agree that the headings used in this Agreement are for their mutual convenience only, and that such headings do not have, nor can be interpreted to have any binding legal effect.",
            "1.3 In case any one or more of the provisions herein contained should be declared invalid, illegal or unenforceable in any respect by any tribunal or court of competent jurisdiction, the validity, legality and enforceability of the remaining provisions hereof shall not in any way be affected or impaired thereby.",
        ]),
        ("2. DEFINITIONS", [
            "2.1 Unless the context otherwise requires, then in construing this Agreement:",
            "2.1.1 Agreement means the entire contractual agreement between the Parties embodied in this Agreement.",
            "2.1.2 Services Materials means, but shall not be limited to, any and all drawings, documents, computer software, information and data stored by any means which is or has been created, directly or indirectly, in connection with or for the purpose of, performing the Services.",
            "2.1.3 Services means the professional services and other services which the Tutor shall provide whilst acting as Tutor as specified in Annex 1 attached hereto which shall form integral part of this Agreement.",
            "2.1.4 Intellectual Property means all copyright, patents, registered and unregistered trademarks, registered designs, trade secrets, and know-how and all other intellectual property as defined in Article 2(vii) of the Convention Establishing the World Intellectual Property Organization of April 1970.",
        ]),
        ("3. PROVISION OF SERVICES AND TERM", [
            "3.1 The Foundation hereby engages the Tutor to provide the Individualised Tuition Services to the Foundation on a professional basis. Nothing in the Agreement shall be construed so as to make the Tutor an employee of the Foundation.",
            "3.2 The Tutor shall exercise reasonable skill, care and diligence in the performance of the Services in accordance with the terms of this Agreement.",
            "3.3 The Tutor shall provide Services to the Foundation exclusively throughout the duration of this Agreement.",
            f"3.4 The term of this Agreement shall run from {start} until such agreement is terminated in accordance with Clause 6 of this Agreement.",
            f"3.5 The Tutor will respond to all communication requests by the Foundation, and its clients and on behalf of the Foundation within a timely manner. The Tutor will represent themselves as part of the team of Learning 360 Foundation and use the following email address {_BLANK}.",
            "3.6 The Tutor shall initially report to Ms. Francesca Diacono, Ms Samantha Pace or Ms Justine Balani and take instructions from either. The Tutor shall remain available to the Foundation on the Tutor's cell phone which the Tutor shall notify to the Foundation and the Tutor agrees that they shall promptly notify the Foundation should such cell phone number change.",
            "3.7 Throughout the term of this Agreement the Tutor shall perform those duties, which are commensurate with their post, including but not limited to those duties specified in Annex 1.",
        ]),
        ("4. KNOWLEDGE OF THE FOUNDATION'S REQUIREMENTS", [
            "4.1 The Tutor shall use all reasonable efforts to remain informed of the Foundation's requirements and for that purpose shall consult the Foundation throughout the performance of the Services.",
        ]),
        ("5. PAYMENT", [
            "5.1 In relation to each service provided by the Tutor, the client will forward payment to Learning360° Foundation.",
            "5.2 The Tutor will be entitled to the following hourly rates, for which they need to insert in their monthly timesheet in order to indicate the specific rate for each level of qualification:",
            f"Euro {_BLANK} per hour for office sessions",
            f"Euro {_BLANK} per hour for home sessions",
            "5.3 The Foundation shall reimburse the Tutor on a monthly basis, payable in arrears by the last Friday of the following month for which the payment is due. In stating this, the Tutor is requested to send a timesheet by the end of each month with all hours worked to the email address: samantha@learning360.org.mt",
            "5.4 Payment will be made to the Tutor's bank whose details are as follows:",
            f"Beneficiary Name: {account_holder}",
            f"Beneficiary Address: {address}",
            f"Bank Name: {_BLANK}",
            f"Bank Branch: {_BLANK}",
            f"Account/IBAN No: {iban}",
        ]),
        ("6. TERMINATION", [
            "6.1 The Foundation shall have the right to terminate this Agreement forthwith by simple written notice to the Tutor in the event of the Tutor's failure to perform the Services up to the standard reasonably expected by the Foundation or in the case of the Tutor's fraud, negligence or willful misconduct.",
            "6.2 Each Party may unilaterally terminate this Agreement without being obliged to attribute a reason for doing so by giving the other party 1 month prior notice in writing.",
            "6.3 Termination of this Agreement shall be without prejudice to those obligations and covenants, including without limitation relating to confidentiality and non-competition, which shall survive the termination of the Agreement.",
        ]),
        ("7. INDEMNITY", [
            "7.1 The Tutor hereby agrees to indemnify on demand the Foundation for and hold the Foundation harmless against any and all costs, losses or damages of whatsoever nature suffered by the Foundation as a result of the Tutor's fraud, negligence or willful misconduct or the Tutor's breach of any provision of this Agreement.",
        ]),
        ("8. PROVISION OF DOCUMENTS ON TERMINATION", [
            "8.1 In addition to, and without in any way diminishing any and all rights of action under contract or at law that the Foundation may have against the Tutor, the Tutor shall, upon termination of the Agreement for any reason whatsoever, immediately surrender to the Foundation all Services Material.",
        ]),
        ("9. SERVICES MATERIAL", [
            "9.1 Subject to any agreement in writing between the Parties to the contrary, the title to and Intellectual Property rights in or in relation to the Services Material shall vest upon its creation exclusively in the Foundation.",
            "9.2 The Tutor shall not use the Services Material for any purposes other than the performance of the Services and shall take all reasonable steps to ensure that no unauthorized person has access to the Services Material during or after the completion of the Services.",
        ]),
        ("10. NON-COMPETITION AND NON-SOLICITATION", [
            "10.1 For the purposes of this Clause 10:",
            "a) \"Foundation Employee\" means any person who is or was employed by the Foundation or any Group Foundation and: i. with whom the Tutor had personal contact or dealings in performing of their duties under this Agreement; or ii. who reported to the Tutor; or iii. who was an employee and/or manager and/or agent and/or director and/or officer of the Foundation.",
            "b) \"Client\" shall mean any person, firm, company or other organization whatsoever to whom the Foundation has supplied goods or services during the five (5) years prior to the date of this Agreement.",
            "c) \"Prospective Client\" shall mean any person, firm, company or other organization whatsoever to whom, during the term of this Agreement, the Foundation has offered to supply goods or services, or has provided details of the terms in which it would or might be willing to supply goods or services, or with whom the Foundation has had any negotiations or discussions regarding the possible supply of goods or services.",
            "d) \"Relevant Period\" shall mean three (3) years from the last day on which the Tutor provides their services to the Foundation in terms of this Agreement.",
            "This Clause is intended to be read, interpreted and enforced in its broadest application and shall have the effect as if it were separate clauses, each being severable from the other. In the event that any of the said provisions shall be declared invalid or unenforceable for any reason by a tribunal or court of competent jurisdiction, such invalidity or unenforceability shall not affect the validity or enforceability of any of the other separate provisions.",
            "10.2 The Tutor hereby agrees that they shall not, during the term of this Agreement and for the Relevant Period, whether on their own behalf or in conjunction with or on behalf of any other person, firm, company, business entity or other organization whatsoever, in competition with the Foundation, directly or indirectly solicit, or assist in soliciting, or accept, or facilitate the acceptance of, or deal with, any Client or Prospective Client.",
            "10.3 The Tutor hereby agrees that they shall not, during the term of this Agreement and for the Relevant Period, either on their own or in conjunction with or on behalf of any other person, firm, company, business entity or other organization whatsoever directly or indirectly: (a) induce, solicit, entice or procure, any person who is a Foundation Employee or Tutor to leave such employment or engagement with the Foundation or; (b) accept into employment or otherwise engage or use the services of any person who is a Foundation Employee.",
            "10.4 The Tutor and the Foundation acknowledge and agree that: a) each of the sub-clauses contained in this Clause constitutes an entirely separate, severable and independent covenant and restriction on the Tutor; and b) the duration, extent and application of each of the restrictions contained in this Agreement are no greater than is necessary for the protection of the goodwill, customer connection, stable workforce and trade connections of the Foundation; and c) in the event that any restriction on the Tutor contained in this Agreement shall be found void but would be valid if some part whereof would be deleted such restriction shall apply with any such deletion as may be necessary to make it valid and effective; and d) in case of any breach of obligations of the Tutor in this clause, the Parties agree that the Tutor shall forfeit any payments due to them under this Agreement.",
            "10.5 The above sub-clauses in this Clause shall also apply as though references to the \"Group Foundation\" were substituted for references to the \"Foundation\". The obligations undertaken by the Tutor pursuant to this Clause shall, with respect to each Group Foundation, constitute a separate and distinct covenant and the invalidity or unenforceability of any such covenant shall not affect the validity and enforceability of the covenants in favor of the Foundation or other Group Foundation provided always that this clause shall only apply to those Group Companies to whom the Tutor rendered their services, or with whom the Tutor was concerned, or in respect of whom the Tutor was responsible.",
        ]),
        ("11. CONFIDENTIALITY", [
            "11.1 For the purposes of this Clause 11, “Confidential Information” shall mean all of the following (whether written, machine-reproducible, oral or in other form or format): Intellectual Property, details of operations, methods, know-how, studies, practices, technical plans, customer lists, customer and prospective customer information and agreements, price lists, supplier lists, supplier information and agreements, marketing plans, financial information, contracts and all other compilations of information which relate to the business or proprietary information of the Foundation and which derive independent economic value from not being generally known to, and not being readily ascertainable by proper means by, other persons and which are the subject of efforts on the part of the Foundation to maintain its secrecy that are reasonable under the circumstances; and other documents into which Confidential Information may be incorporated. The term “Confidential Information” does not include information or material that (i) was in the public domain prior to the date of this Agreement or that subsequently comes into the public domain through no fault of the Tutor; or (ii) is required to be disclosed in a judicial or administrative proceeding or by a governmental or regulatory authority provided that the Tutor shall promptly advise the Foundation in writing upon becoming aware that any such disclosure may be imminent.",
            "11.2 The Tutor shall: a) use utmost care to maintain the confidentiality of Confidential Information and restrict and protect its disclosure; b) not copy Confidential Information for any purpose other than to perform the Services (such copies shall for the purposes of this Agreement be considered to be Services Material); c) not discuss or disclose Confidential Information except as authorized by this Agreement or as expressly authorized by the Foundation; and d) promptly report to the Foundation any actual or suspected violation of the terms of this Agreement and take all reasonable further steps required by the Foundation to prevent, mitigate, control or remedy any such violation.",
        ]),
        ("12. SUCCESSORS, BINDING AGREEMENT", [
            "12.1 The Foundation cannot assign this Agreement, without written consent of the Tutor. This Agreement shall inure to the benefit of the Tutor and their personal or legal representative, executors, estate, trustee, administrators, successors, heirs, distributees, devisees and legatees. The Tutor may not assign this Agreement or any rights hereunder. However, in the event of the Tutor's death, all rights to receive payments accrued by the Tutor until the time of their death, shall become rights of the Tutor's devisees, legatees or other designees or the Tutor's estate.",
        ]),
        ("13. AUTHORITY TO CONTRACT", [
            "13.1 The Tutor represents and warrants that they have the full right, power and authority to enter into this Agreement and that this Agreement constitutes a valid obligation of the Tutor, binding on the Tutor and enforceable in accordance with the terms hereof. In addition, the Tutor represents and warrants that the execution of this Agreement and the performance of their obligations herein will not conflict with or violate any commitment, agreement, contract or understanding the Tutor has or will have with any other person or entity and there is nothing that will prevent the Tutor from performing their obligations under the terms and conditions imposed on them by this Agreement.",
        ]),
        ("14. NO WAIVER", [
            "14.1 Each Party agrees that no failure or delay by the other Party in exercising any right, power or privilege hereunder will operate as a waiver thereof, nor will any single or partial exercise thereof preclude any other or further exercise thereof or the exercise of any right, power or privilege hereunder.",
        ]),
        ("15. GOVERNING LAW AND JURISDICTION", [
            "15.1 This Agreement shall be governed by and construed in accordance with the laws of Malta, and for the benefit of the Foundation, any dispute in respect of this Agreement shall be submitted to the exclusive jurisdiction of the Courts of Malta.",
        ]),
        ("16. COUNTERPARTS", [
            "16.1 This Agreement may be executed in several counterparts (including through signed pdf copies exchanged via email), each of which shall be deemed to be the original but all of which together will constitute one and the same instrument.",
        ]),
    ]

    for heading, paras in sections:
        _para(doc, heading, bold=True, size=11)
        for t in paras:
            _para(doc, t)

    _para(doc, "Thus agreed and signed on the date first aforementioned:", bold=True)
    _para(doc, "SIGNED for and on behalf of Learning 360 Foundation:")
    for founder in ("Francesca Diacono", "Samantha Pace", "Justine Balani"):
        _para(doc, f"By ………………………………  {founder}")
    _para(doc, f"SIGNED by the Tutor:")
    _para(doc, f"By ………………………………  {name}")

    _para(doc, "ANNEX 1 — The Services", bold=True, size=12)
    for t in (
        "Throughout the term of this Agreement the Tutor shall provide the following Services to the Foundation and cooperate as necessary in the best interest of the Foundation and its businesses with any other Tutor/s appointed by the Foundation to provide similar or ancillary services. The Foundation shall have the right to request the Tutor to give priority to certain services during any particular period or periods. It is understood that this list is subject to change based on time and budgetary constraints.",
        "Policies and Procedures — Duties and responsibilities of a Tutor.",
        "Learning360° Foundation provides holistic educational services, specifically: support, coaching, training and mentoring for students, educators, parents and professionals. The aim is to offer specialized, individualized and group educational services addressing skills and holistic development such as early intervention, study skills, social skills and mentoring among others.",
        "Learning360° Foundation Tutors are required to fill in an initial consultation form upon the first session. This gives the parents the opportunity to provide the Tutor with lots of background information about their child, their strengths, areas of improvement, targets and goals they wish to work on. It gives the Tutor the opportunity to start to build a relationship with the parents and the student and to understand both the parents' needs and wishes. Ongoing tracking of goals is required in order to facilitate growth and achievement.",
        "At Learning360° Foundation our duty as a Tutor is to offer individualised support according to that child's individual needs and profile. Each family differs and so do their expectations; we strive to provide excellent tutor services and to be as accommodating as we can.",
        "Duties of a Tutor shall include, but are not limited to:",
        "• Participates and collaborates in all meetings held with guardians and other stakeholders as per the individual needs of the client.",
        "• Is responsible for the planning, development and implementation of the individual educational plan of the client, in consultation with guardians, other stakeholders and supervisors as needed.",
        "• Tracks strategies and goals on a weekly basis in the provided log sheet and consults with supervisor/s on any concerns/queries should the targeted goals not be on track or no progress is detected, given an appropriate time frame.",
        "• Prepares and adapts resources and materials needed for the individual programme of the child where necessary, including caring for and returning in good condition any resources which are borrowed from Learning360° Foundation.",
        "• Abides by the most recent policies and procedures of Learning360° Foundation.",
        "• Allows two to three sessions for a handover to another tutor, in case of resignation during the set notice period of four weeks.",
        "• Should the Tutor need to cancel/postpone a session with a client, they are to inform the client and Learning360° Foundation at least within 24 hours of the scheduled appointment, unless a cancellation is due to sickness/unforeseeable circumstances.",
        "• Participates and collaborates in opportunities for professional development and/or training and mentoring as per the advice of the supervisor.",
        "• Promotes at all times the aims, ethos and policies of Learning360° Foundation.",
        "• Appropriate clothing will be worn during all sessions. Dress code should be smart casual; short, tight-fitting, see-through, low-cleavage clothing should be avoided.",
        "Code of Behaviour",
        "• Every student in my care will be treated with respect, and I will regard their well-being and safety.",
        "• I will forge professional and supportive relationships with family members as well as other stakeholders.",
        "Confidentiality",
        "• I respect the privacy of the families I work with. No details about any student are ever discussed with anyone other than the child's parents/guardians.",
    ):
        _para(doc, t)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
